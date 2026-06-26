from __future__ import annotations

import csv
import json
import os
import re
import secrets
import shutil
import sys
import threading
import uuid
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from fastapi import Depends, FastAPI, File, Form, Header, HTTPException, Request, UploadFile
from fastapi.concurrency import run_in_threadpool
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, PlainTextResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel


APP_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = APP_DIR.parent
SRC_DIR = PROJECT_ROOT / "src"
RUNS_DIR = APP_DIR / "runs"
SITE_DIR = APP_DIR / "site"

if str(SRC_DIR) not in sys.path:
    sys.path.insert(0, str(SRC_DIR))

from chunk_book import chunk_book_text, detect_chapter_spans, word_count  # noqa: E402
from eval_external_outputs import (  # noqa: E402
    QUESTION_ID_RE,
    _answer_found_for_summary,
    _load_questions,
    _qid_segments,
    _score_question,
)
from run_eval import METHODS as EVAL_METHODS, run_eval as run_eval_engine  # noqa: E402


app = FastAPI(
    title="LongBook Verifier MVP",
    description="Verify-only local MVP for grounding AI outputs against uploaded long manuscripts.",
    version="0.1.0",
)

if SITE_DIR.exists():
    app.mount("/site", StaticFiles(directory=str(SITE_DIR)), name="site")

# The BookProof browser app calls this API. In production it is same-origin via the
# Netlify /api/* proxy; these origins also allow local browser testing.
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:8000",
        "http://127.0.0.1:8000",
        "http://localhost:8888",
        "http://127.0.0.1:8888",
        "http://localhost:5500",
        "http://127.0.0.1:5500",
        "https://tts.bedvibe.studio",
        "https://bedvibe.studio",
    ],
    allow_methods=["GET", "POST"],
    allow_headers=["*"],
)

BOOKPROOF_DEFAULT_METHOD = "flat_chunk_rag"
BOOKPROOF_BOOK_SUFFIXES = {".txt", ".md", ".docx"}

# Shared-token gate for the compute endpoint. Token comes from the BOOKPROOF_TOKEN
# environment variable (set via the systemd EnvironmentFile). Fails closed: if no
# token is configured, the run endpoint refuses rather than running open.
BOOKPROOF_TOKEN = os.environ.get("BOOKPROOF_TOKEN", "").strip()


def _require_bookproof_token(
    x_bookproof_token: Optional[str] = Header(default=None, alias="X-BookProof-Token"),
) -> None:
    if not BOOKPROOF_TOKEN:
        raise HTTPException(status_code=503, detail="BookProof run gate is not configured")
    if not x_bookproof_token or not secrets.compare_digest(x_bookproof_token, BOOKPROOF_TOKEN):
        raise HTTPException(status_code=401, detail="Missing or invalid X-BookProof-Token")


# One active engine run at a time across all BookProof run endpoints. The engine is
# CPU-bound and shares the box with other services, so concurrent runs are refused.
V1_MAX_TEXT_CHARS = 3_000_000
V1_MAX_CLAIMS = 200
_RUN_LOCK = threading.Lock()


def _single_run_slot():
    if not _RUN_LOCK.acquire(blocking=False):
        raise HTTPException(status_code=429, detail="BookProof is busy, try again shortly.")
    try:
        yield
    finally:
        _RUN_LOCK.release()


# Public free-demo limits (separate, stricter than the token-gated developer API).
PUBLIC_MAX_TEXT_CHARS = 200_000
PUBLIC_MAX_CLAIMS = 10
PUBLIC_DAILY_LIMIT = 1  # validated runs per client IP per UTC day
_public_quota_lock = threading.Lock()
_public_ip_runs: Dict[str, Any] = {}


def _client_ip(request: Request) -> str:
    # Behind Cloudflare + nginx the real visitor IP is in CF-Connecting-IP
    # (X-Forwarded-For first hop as fallback); request.client is the proxy hop.
    cf = request.headers.get("cf-connecting-ip")
    if cf:
        return cf.strip()
    xff = request.headers.get("x-forwarded-for")
    if xff:
        return xff.split(",")[0].strip()
    return request.client.host if request.client else "unknown"


def _enforce_public_quota(ip: str) -> None:
    day = datetime.now(timezone.utc).strftime("%Y%m%d")
    with _public_quota_lock:
        d, n = _public_ip_runs.get(ip, (day, 0))
        if d != day:
            d, n = day, 0
        if n >= PUBLIC_DAILY_LIMIT:
            raise HTTPException(
                status_code=429,
                detail=f"Free public demo limit reached ({PUBLIC_DAILY_LIMIT} run/day per IP). "
                "Try again tomorrow, or use the token-gated developer API for higher limits.",
            )
        _public_ip_runs[ip] = (day, n + 1)


WORD_RE = re.compile(r"[A-Za-z][A-Za-z0-9_'\-]{2,}")
NAME_RE = re.compile(r"\b(?:[A-Z][A-Za-z'’\-]+)(?:\s+[A-Z][A-Za-z'’\-]+){0,3}\b")
SAFE_RUN_ID_RE = re.compile(r"^[A-Za-z0-9_-]+$")

STOPWORDS = {
    "about",
    "above",
    "after",
    "again",
    "against",
    "answer",
    "because",
    "before",
    "being",
    "book",
    "chapter",
    "could",
    "does",
    "from",
    "have",
    "into",
    "manuscript",
    "only",
    "question",
    "result",
    "section",
    "that",
    "their",
    "them",
    "then",
    "there",
    "these",
    "they",
    "this",
    "through",
    "what",
    "when",
    "where",
    "which",
    "while",
    "with",
    "would",
}


@app.get("/")
def site_index() -> FileResponse:
    index = SITE_DIR / "index.html"
    if not index.exists():
        raise HTTPException(status_code=404, detail="MVP site not found")
    return FileResponse(index)


@app.get("/health")
def health() -> Dict[str, bool]:
    return {"ok": True}


def _safe_filename(name: Optional[str], fallback: str) -> str:
    raw = Path(name or fallback).name.strip()
    raw = re.sub(r"[^A-Za-z0-9_.\- ]+", "_", raw)
    return raw or fallback


async def _save_upload(upload: UploadFile, run_dir: Path, label: str, allowed_suffixes: set[str]) -> Path:
    filename = _safe_filename(upload.filename, label)
    suffix = Path(filename).suffix.lower()
    if suffix not in allowed_suffixes:
        allowed = ", ".join(sorted(allowed_suffixes))
        raise HTTPException(status_code=400, detail=f"{label} must use one of: {allowed}")

    path = run_dir / filename
    path.write_bytes(await upload.read())
    return path


def _read_docx_text(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise HTTPException(status_code=500, detail="DOCX support requires python-docx") from exc

    document = Document(str(path))
    parts: List[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts).strip()


def _read_text_file(path: Path, role: str) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".jsonl"}:
        return path.read_text(encoding="utf-8-sig", errors="replace")
    if suffix == ".docx" and role == "book":
        return _read_docx_text(path)
    raise HTTPException(status_code=400, detail=f"Unsupported {role} file type: {suffix}")


def _normalize(text: str) -> str:
    return " ".join(re.findall(r"[A-Za-z0-9_']+", str(text or "").lower()))


def _contains_phrase(normalized_text: str, phrase: str) -> bool:
    phrase_norm = _normalize(phrase)
    return bool(phrase_norm and phrase_norm in normalized_text)


def _top_named_terms(text: str, limit: int = 25) -> List[Dict[str, Any]]:
    counts: Counter[str] = Counter()
    for match in NAME_RE.findall(text):
        term = " ".join(match.split())
        first = term.split()[0].lower()
        if first in STOPWORDS or len(term) < 3:
            continue
        counts[term] += 1
    return [{"term": term, "count": count} for term, count in counts.most_common(limit)]


def _candidate_terms(ai_text: str, limit: int = 250) -> List[str]:
    terms: List[str] = []
    for row in _top_named_terms(ai_text, limit=80):
        term = str(row["term"])
        if term not in terms:
            terms.append(term)

    token_counts: Counter[str] = Counter()
    for token in WORD_RE.findall(ai_text):
        cleaned = token.lower().strip("'_-")
        if len(cleaned) < 5 or cleaned in STOPWORDS:
            continue
        token_counts[cleaned] += 1

    for token, _count in token_counts.most_common(limit):
        if token not in terms:
            terms.append(token)
        if len(terms) >= limit:
            break
    return terms


def _evidence_overlap(book_text: str, ai_text: str) -> Tuple[float, List[str], List[str]]:
    candidates = _candidate_terms(ai_text)
    if not candidates:
        return 0.0, [], []

    book_norm = _normalize(book_text)
    supported = [term for term in candidates if _contains_phrase(book_norm, term)]
    unsupported = [term for term in candidates if term not in supported]
    return len(supported) / float(len(candidates)), supported, unsupported


def _score_questions(ai_text: str, questions_path: Path, csv_path: Path) -> Dict[str, Any]:
    questions = _load_questions(questions_path)
    segments = _qid_segments(ai_text)
    fallback_full_document = not bool(segments or QUESTION_ID_RE.search(ai_text))

    rows = [
        _score_question(
            "uploaded_ai_output",
            ai_text,
            question,
            segments,
            fallback_full_document=fallback_full_document,
        )
        for question in questions
    ]

    columns = [
        "system_name",
        "question_id",
        "qid_found",
        "answer_found",
        "gold_answer_coverage",
        "evidence_coverage",
        "matched_evidence_terms",
        "missing_evidence_terms",
        "expected_chapter",
        "notes",
    ]
    with csv_path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=columns)
        writer.writeheader()
        for row in rows:
            writer.writerow({key: row.get(key, "") for key in columns})

    total = len(rows)
    qids_found = sum(int(row["qid_found"]) for row in rows)
    answers_found = sum(_answer_found_for_summary(row["answer_found"]) for row in rows)
    evidence = sum(float(row["evidence_coverage"]) for row in rows) / float(total or 1)
    answer_terms = sum(float(row["gold_answer_coverage"]) for row in rows) / float(total or 1)

    return {
        "questions": total,
        "qids_found": qids_found,
        "answers_found": answers_found,
        "evidence_coverage": round(evidence, 6),
        "answer_term_coverage": round(answer_terms, 6),
        "csv_rows": rows,
        "fallback_full_document": fallback_full_document,
    }


def _write_report(
    report_path: Path,
    run_id: str,
    book_path: Path,
    ai_path: Path,
    book_words: int,
    token_estimate: int,
    chapter_count: int,
    chunk_count: int,
    top_terms: List[Dict[str, Any]],
    evidence_coverage: float,
    unsupported_terms: List[str],
    question_summary: Optional[Dict[str, Any]],
    csv_path: Optional[Path],
) -> None:
    lines = [
        "# LongBook Verifier MVP Report",
        "",
        f"- Run ID: `{run_id}`",
        f"- Book upload: `{book_path.name}`",
        f"- AI output upload: `{ai_path.name}`",
        f"- Book word count: {book_words:,}",
        f"- Approximate token count: {token_estimate:,}",
        f"- Detected chapters/sections: {chapter_count}",
        f"- Generated chunks: {chunk_count}",
        "",
        "## Scores",
        "",
        f"- Evidence overlap / evidence coverage: {evidence_coverage:.4f}",
    ]
    if question_summary:
        lines.extend(
            [
                f"- Answer-term coverage: {question_summary['answer_term_coverage']:.4f}",
                f"- QIDs found: {question_summary['qids_found']} / {question_summary['questions']}",
                f"- Answers found: {question_summary['answers_found']} / {question_summary['questions']}",
                f"- Fallback full-document scoring used: {question_summary['fallback_full_document']}",
            ]
        )
        if csv_path:
            lines.append(f"- Question score CSV: `{csv_path.name}`")

    lines.extend(["", "## Top Repeated Named Terms", ""])
    if top_terms:
        for row in top_terms[:20]:
            lines.append(f"- {row['term']} ({row['count']})")
    else:
        lines.append("- No repeated named terms detected.")

    lines.extend(["", "## Possible Unsupported Terms In AI Output", ""])
    if unsupported_terms:
        for term in unsupported_terms[:40]:
            lines.append(f"- {term}")
    else:
        lines.append("- No unsupported term candidates detected by the lightweight overlap check.")

    lines.extend(
        [
            "",
            "## Notes",
            "",
            "This MVP is verify-only. It does not generate summaries or answers, and it does not call paid model APIs.",
            "Automatic evidence overlap is lexical and lightweight; question JSONL scoring is stronger because it uses expected evidence_terms and gold_answer fields.",
        ]
    )
    report_path.write_text("\n".join(lines) + "\n", encoding="utf-8")


@app.post("/api/verify")
async def verify(
    book_file: UploadFile = File(...),
    ai_output_file: UploadFile = File(...),
    questions_file: Optional[UploadFile] = File(None),
) -> Dict[str, Any]:
    run_id = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S_") + uuid.uuid4().hex[:8]
    run_dir = RUNS_DIR / run_id
    run_dir.mkdir(parents=True, exist_ok=True)

    book_path = await _save_upload(book_file, run_dir, "book_file", {".txt", ".md", ".docx"})
    ai_path = await _save_upload(ai_output_file, run_dir, "ai_output_file", {".txt", ".md"})

    questions_path: Optional[Path] = None
    if questions_file is not None and questions_file.filename:
        questions_path = await _save_upload(questions_file, run_dir, "questions_file", {".jsonl"})

    book_text = _read_text_file(book_path, "book")
    ai_text = _read_text_file(ai_path, "ai output")

    book_words = word_count(book_text)
    token_estimate = int(round(book_words * 1.33))
    chapter_count = len(detect_chapter_spans(book_text))
    chunks = chunk_book_text(book_text)
    top_terms = _top_named_terms(book_text)
    auto_evidence_coverage, _supported_terms, unsupported_terms = _evidence_overlap(book_text, ai_text)

    csv_path: Optional[Path] = None
    question_summary: Optional[Dict[str, Any]] = None
    evidence_coverage = round(auto_evidence_coverage, 6)
    answer_term_coverage: Optional[float] = None
    qids_found: Optional[int] = None
    answers_found: Optional[int] = None

    if questions_path is not None:
        csv_path = run_dir / "question_scores.csv"
        question_summary = _score_questions(ai_text, questions_path, csv_path)
        evidence_coverage = float(question_summary["evidence_coverage"])
        answer_term_coverage = float(question_summary["answer_term_coverage"])
        qids_found = int(question_summary["qids_found"])
        answers_found = int(question_summary["answers_found"])

    report_path = run_dir / "report.md"
    _write_report(
        report_path=report_path,
        run_id=run_id,
        book_path=book_path,
        ai_path=ai_path,
        book_words=book_words,
        token_estimate=token_estimate,
        chapter_count=chapter_count,
        chunk_count=len(chunks),
        top_terms=top_terms,
        evidence_coverage=evidence_coverage,
        unsupported_terms=unsupported_terms,
        question_summary=question_summary,
        csv_path=csv_path,
    )

    return {
        "run_id": run_id,
        "evidence_coverage": evidence_coverage,
        "answer_term_coverage": answer_term_coverage,
        "qids_found": qids_found,
        "answers_found": answers_found,
        "unsupported_terms_preview": unsupported_terms[:25],
        "report_path": str(report_path),
        "csv_path": str(csv_path) if csv_path else None,
    }


@app.get("/api/report/{run_id}")
def get_report(run_id: str) -> PlainTextResponse:
    if not SAFE_RUN_ID_RE.fullmatch(run_id):
        raise HTTPException(status_code=400, detail="Invalid run_id")

    report_path = RUNS_DIR / run_id / "report.md"
    if not report_path.exists():
        raise HTTPException(status_code=404, detail="Report not found")
    return PlainTextResponse(report_path.read_text(encoding="utf-8"), media_type="text/markdown")


# ---------------------------------------------------------------------------
# BookProof bridge: browser -> this backend API -> existing LongBook retrieval
# evaluation engine. Retrieval/evaluation only, deterministic (hashing_numpy).
# No external model API calls, and no BYOK key is ever read, stored, or logged.
# ---------------------------------------------------------------------------


@app.get("/api/bookproof/health")
def bookproof_health() -> Dict[str, Any]:
    return {
        "ok": True,
        "service": "bookproof",
        "engine": "longbook-retrieval-eval",
        "embedding_backend": "hashing_numpy",
        "model_calls": False,
        "methods": sorted(EVAL_METHODS),
        "runs_dir": str(RUNS_DIR),
    }


def _validate_questions_jsonl(text: str) -> None:
    saw = False
    for line_no, raw in enumerate(text.splitlines(), start=1):
        line = raw.strip()
        if not line:
            continue
        try:
            obj = json.loads(line)
        except json.JSONDecodeError as exc:
            raise HTTPException(
                status_code=400,
                detail=f"Golden questions line {line_no} is not valid JSON: {exc.msg}",
            ) from exc
        if not isinstance(obj, dict) or not str(obj.get("question") or "").strip():
            raise HTTPException(
                status_code=400,
                detail=f"Golden questions line {line_no} must be a JSON object with a non-empty 'question' field",
            )
        saw = True
    if not saw:
        raise HTTPException(status_code=400, detail="No golden questions provided")


def _write_bookproof_report(
    run_dir: Path, run_id: str, summary: Dict[str, Any], results: List[Dict[str, Any]]
) -> Path:
    def f4(value: Any) -> str:
        try:
            return f"{float(value):.4f}"
        except (TypeError, ValueError):
            return "—"

    lines = [
        "# BookProof Retrieval Evaluation Report",
        "",
        f"- Run ID: `{run_id}`",
        f"- Method: `{summary.get('method')}` ({summary.get('method_description', '')})",
        f"- Embedding backend: `{summary.get('embedding_backend')}` (deterministic, no model API)",
        f"- Questions evaluated: {summary.get('count')}",
        f"- Book chunks: {summary.get('chunk_count')}  ·  chapters/sections: {summary.get('chapter_count')}",
        f"- top_k: {summary.get('top_k')}",
        "",
        "## Aggregate scores",
        "",
        f"- Avg context precision (like): {f4(summary.get('avg_context_precision'))}",
        f"- Avg context recall (like): {f4(summary.get('avg_context_recall'))}",
        f"- Avg answer-term coverage: {f4(summary.get('avg_answer_score'))}",
        f"- Avg retrieval latency (s): {f4(summary.get('avg_latency'))}",
        "",
        "## Per-question (first 20)",
        "",
    ]
    if results:
        for row in results[:20]:
            question = " ".join(str(row.get("question") or "").split())[:160]
            metrics = row.get("metrics") or {}
            lines.append(
                f"- `{row.get('id')}` — precision {f4(metrics.get('context_precision_like'))}, "
                f"recall {f4(metrics.get('context_recall_like'))}, "
                f"answer {f4(metrics.get('answer_contains_gold_terms'))} — {question}"
            )
    else:
        lines.append("- (no questions scored)")
    lines.extend(
        [
            "",
            "## Notes",
            "",
            "Retrieval-evaluation only. Deterministic local embeddings (hashing_numpy); no Claude/OpenAI/Gemini calls "
            "and no model answer generation. Run artifacts in this folder: summary.json, metrics.csv, results.jsonl.",
        ]
    )
    report_path = run_dir / "report.md"
    report_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return report_path


@app.post("/api/bookproof/run", dependencies=[Depends(_require_bookproof_token), Depends(_single_run_slot)])
async def bookproof_run(
    book_file: UploadFile = File(...),
    questions_file: Optional[UploadFile] = File(None),
    questions_text: Optional[str] = Form(None),
    method: str = Form(BOOKPROOF_DEFAULT_METHOD),
    top_k: int = Form(5),
    provider: Optional[str] = Form(None),  # reserved/unused — this engine makes no model calls
    model: Optional[str] = Form(None),     # reserved/unused
) -> Dict[str, Any]:
    # No `api_key` parameter is declared, so a BYOK key is never read, stored, or logged here.
    if method not in EVAL_METHODS:
        raise HTTPException(
            status_code=400, detail=f"method must be one of: {', '.join(sorted(EVAL_METHODS))}"
        )
    top_k = max(1, min(int(top_k), 20))

    run_id = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S_") + uuid.uuid4().hex[:8]
    run_dir = RUNS_DIR / run_id
    run_dir.mkdir(parents=True, exist_ok=True)

    book_path = await _save_upload(book_file, run_dir, "book_file", BOOKPROOF_BOOK_SUFFIXES)

    # Golden questions: a .jsonl upload OR pasted JSONL text. Always written to a fixed,
    # collision-free name inside the run folder only.
    questions_path = run_dir / "input_questions.jsonl"
    if questions_file is not None and questions_file.filename:
        if Path(_safe_filename(questions_file.filename, "questions.jsonl")).suffix.lower() != ".jsonl":
            raise HTTPException(status_code=400, detail="questions_file must be a .jsonl file")
        raw = (await questions_file.read()).decode("utf-8-sig", errors="replace")
        _validate_questions_jsonl(raw)
        questions_path.write_text(raw.strip() + "\n", encoding="utf-8")
    elif questions_text and questions_text.strip():
        _validate_questions_jsonl(questions_text)
        questions_path.write_text(questions_text.strip() + "\n", encoding="utf-8")
    else:
        raise HTTPException(
            status_code=400, detail="Provide golden questions as a .jsonl file or as JSONL text"
        )

    # Run the existing deterministic engine in a worker thread (CPU-bound).
    summary = await run_in_threadpool(
        run_eval_engine,
        book=book_path,
        questions_path=questions_path,
        method=method,
        out=run_dir,
        top_k=top_k,
        embedding_backend="hashing_numpy",
    )

    results: List[Dict[str, Any]] = []
    results_file = run_dir / "results.jsonl"
    if results_file.exists():
        for raw in results_file.read_text(encoding="utf-8").splitlines():
            raw = raw.strip()
            if not raw:
                continue
            try:
                results.append(json.loads(raw))
            except json.JSONDecodeError:
                continue

    report_path = _write_bookproof_report(run_dir, run_id, summary, results)

    preview = [
        {
            "id": row.get("id"),
            "question": " ".join(str(row.get("question") or "").split())[:200],
            "selected_chapters": row.get("selected_chapters"),
            "context_precision_like": (row.get("metrics") or {}).get("context_precision_like"),
            "context_recall_like": (row.get("metrics") or {}).get("context_recall_like"),
            "answer_contains_gold_terms": (row.get("metrics") or {}).get("answer_contains_gold_terms"),
        }
        for row in results[:10]
    ]

    return {
        "run_id": run_id,
        "status": "ok",
        "method": summary.get("method"),
        "embedding_backend": summary.get("embedding_backend"),
        "questions_evaluated": summary.get("count"),
        "chunk_count": summary.get("chunk_count"),
        "chapter_count": summary.get("chapter_count"),
        "top_k": summary.get("top_k"),
        "scores": {
            "avg_context_precision": summary.get("avg_context_precision"),
            "avg_context_recall": summary.get("avg_context_recall"),
            "avg_answer_term_coverage": summary.get("avg_answer_score"),
            "avg_latency_seconds": summary.get("avg_latency"),
        },
        "results_preview": preview,
        "report_url": f"/api/bookproof/report/{run_id}",
        "report_path": str(report_path),
        "artifacts": {
            "summary_json": str(run_dir / "summary.json"),
            "metrics_csv": str(run_dir / "metrics.csv"),
            "results_jsonl": str(results_file),
        },
        "reserved": {
            "provider": provider,
            "model": model,
            "note": "Model answer generation (BYOK) is not enabled in this build.",
        },
    }


@app.get("/api/bookproof/report/{run_id}")
def bookproof_report(run_id: str) -> PlainTextResponse:
    if not SAFE_RUN_ID_RE.fullmatch(run_id):
        raise HTTPException(status_code=400, detail="Invalid run_id")
    report_path = RUNS_DIR / run_id / "report.md"
    if not report_path.exists():
        raise HTTPException(status_code=404, detail="Report not found")
    return PlainTextResponse(report_path.read_text(encoding="utf-8"), media_type="text/markdown")


# ---------------------------------------------------------------------------
# API v1 — agent/script callable. Send document text + claims/questions as JSON,
# no file upload and no browser. Token-gated, one-run-at-a-time, size-limited.
# Deterministic local retrieval only: no URL fetching, no external model calls.
# ---------------------------------------------------------------------------


class V1VerifyRequest(BaseModel):
    text: str
    claims: List[Any]
    title: Optional[str] = None
    method: str = BOOKPROOF_DEFAULT_METHOD
    top_k: int = 5


def _normalize_v1_claims(raw_claims: List[Any]) -> List[Dict[str, Any]]:
    """Accept both simple string claims and rich claim objects; return engine question rows.
    Rich object: {claim, gold_answer?, evidence_terms?, expected_chapters?}. Backwards compatible."""
    rows: List[Dict[str, Any]] = []
    for idx, item in enumerate(raw_claims, start=1):
        cid = f"c{idx:03d}"
        if isinstance(item, str):
            q = item.strip()
            if q:
                rows.append({"id": cid, "question": q})
            continue
        if isinstance(item, dict):
            q = str(item.get("claim") or item.get("question") or "").strip()
            if not q:
                raise HTTPException(status_code=400, detail=f"claim #{idx} object is missing 'claim' text")
            row: Dict[str, Any] = {"id": cid, "question": q}
            gold = item.get("gold_answer")
            if isinstance(gold, str) and gold.strip():
                row["gold_answer"] = gold.strip()
            terms = item.get("evidence_terms")
            if isinstance(terms, list):
                clean = [str(t).strip() for t in terms if str(t).strip()]
                if clean:
                    row["evidence_terms"] = clean
            chapters = item.get("expected_chapters")
            if isinstance(chapters, list):
                clean_ch = [str(c).strip() for c in chapters if str(c).strip()]
                if clean_ch:
                    row["expected_chapter"] = ", ".join(clean_ch)
                    row["expected_chapters"] = clean_ch
            rows.append(row)
            continue
        raise HTTPException(status_code=400, detail=f"claim #{idx} must be a string or a claim object")
    if not rows:
        raise HTTPException(status_code=400, detail="'claims' must contain at least one non-empty claim")
    return rows


@app.get("/api/bookproof/v1/spec")
def bookproof_v1_spec() -> Dict[str, Any]:
    """Public, machine-readable description of how to call the BookProof API."""
    return {
        "name": "BookProof API",
        "version": "v1",
        "base_url": "https://bookproof.bedvibe.studio",
        "description": "Document-grounded verification and retrieval-evaluation. Send a document and claims/questions; BookProof scores how well each is grounded in the document. Deterministic local retrieval only — no external model APIs.",
        "auth": {"header": "X-BookProof-Token", "required_for": ["POST /api/bookproof/v1/verify"]},
        "endpoints": {
            "GET /api/bookproof/health": "Public health check.",
            "GET /api/bookproof/v1/spec": "This document (public).",
            "POST /api/bookproof/v1/verify": "Verify claims/questions against a document (token required).",
            "GET /api/bookproof/report/{run_id}": "Markdown report for a run.",
        },
        "verify_request_schema": {
            "title": "string (optional) — document title",
            "text": "string (required) — the source document text",
            "claims": "array (required) — each item is EITHER a string (claim/question) OR a claim object (see claim_object_schema). The two forms may be mixed in one request.",
            "claim_object_schema": {
                "claim": "string (required) — claim or question text",
                "gold_answer": "string (optional) — expected answer; enables answer-term coverage scoring",
                "evidence_terms": "string[] (optional) — terms expected in supporting evidence; enables context precision/recall scoring",
                "expected_chapters": "string[] (optional) — expected chapter/section names (carried into results)",
            },
            "method": "string (optional, default '" + BOOKPROOF_DEFAULT_METHOD + "'); one of " + str(sorted(EVAL_METHODS)),
            "top_k": "integer (optional, default 5, clamped 1-20)",
        },
        "limits": {
            "max_text_chars": V1_MAX_TEXT_CHARS,
            "max_claims": V1_MAX_CLAIMS,
            "concurrency": "one active run at a time; returns HTTP 429 when busy",
        },
        "constraints": [
            "No external model API calls (no OpenAI / Claude / Gemini).",
            "No URL fetching — send text directly.",
            "Deterministic embeddings (hashing_numpy).",
        ],
        "example_curl": (
            "curl -s -X POST https://bookproof.bedvibe.studio/api/bookproof/v1/verify "
            "-H 'X-BookProof-Token: <token>' -H 'Content-Type: application/json' "
            "-d '{\"title\":\"My doc\",\"text\":\"...document text...\","
            "\"claims\":[\"a simple claim string\","
            "{\"claim\":\"a richer claim\",\"gold_answer\":\"expected answer\","
            "\"evidence_terms\":[\"term1\",\"term2\"],\"expected_chapters\":[\"ch01\"]}],"
            "\"method\":\"flat_chunk_rag\",\"top_k\":5}'"
        ),
    }


@app.post(
    "/api/bookproof/v1/verify",
    dependencies=[Depends(_require_bookproof_token), Depends(_single_run_slot)],
)
async def bookproof_v1_verify(req: V1VerifyRequest) -> Dict[str, Any]:
    text = (req.text or "").strip()
    if not text:
        raise HTTPException(status_code=400, detail="'text' is required")
    if len(text) > V1_MAX_TEXT_CHARS:
        raise HTTPException(status_code=413, detail=f"'text' exceeds max of {V1_MAX_TEXT_CHARS} characters")

    raw_claims = req.claims or []
    if not isinstance(raw_claims, list) or not raw_claims:
        raise HTTPException(status_code=400, detail="'claims' must be a non-empty list")
    if len(raw_claims) > V1_MAX_CLAIMS:
        raise HTTPException(status_code=400, detail=f"too many claims; max is {V1_MAX_CLAIMS}")
    question_rows = _normalize_v1_claims(raw_claims)

    method = (req.method or BOOKPROOF_DEFAULT_METHOD)
    if method not in EVAL_METHODS:
        raise HTTPException(status_code=400, detail=f"method must be one of: {', '.join(sorted(EVAL_METHODS))}")
    top_k = max(1, min(int(req.top_k or 5), 20))

    run_id = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S_") + uuid.uuid4().hex[:8]
    run_dir = RUNS_DIR / run_id
    run_dir.mkdir(parents=True, exist_ok=True)

    # Temporary input lives ONLY under this run folder.
    doc_path = run_dir / "input_document.txt"
    doc_path.write_text(text, encoding="utf-8")
    questions_path = run_dir / "input_questions.jsonl"
    with questions_path.open("w", encoding="utf-8") as fh:
        for row in question_rows:
            fh.write(json.dumps(row, ensure_ascii=False) + "\n")

    summary = await run_in_threadpool(
        run_eval_engine,
        book=doc_path,
        questions_path=questions_path,
        method=method,
        out=run_dir,
        top_k=top_k,
        embedding_backend="hashing_numpy",
    )

    results: List[Dict[str, Any]] = []
    results_file = run_dir / "results.jsonl"
    if results_file.exists():
        for raw in results_file.read_text(encoding="utf-8").splitlines():
            raw = raw.strip()
            if not raw:
                continue
            try:
                results.append(json.loads(raw))
            except json.JSONDecodeError:
                continue

    _write_bookproof_report(run_dir, run_id, summary, results)

    evidence = [
        {
            "id": row.get("id"),
            "claim": " ".join(str(row.get("question") or "").split())[:500],
            "selected_chapters": row.get("selected_chapters"),
            "context_precision_like": (row.get("metrics") or {}).get("context_precision_like"),
            "context_recall_like": (row.get("metrics") or {}).get("context_recall_like"),
            "answer_contains_gold_terms": (row.get("metrics") or {}).get("answer_contains_gold_terms"),
            "retrieved_chunks": [
                {"id": c.get("id"), "score": c.get("score"), "chapter_id": c.get("chapter_id")}
                for c in (row.get("retrieved_chunks") or [])
            ],
        }
        for row in results
    ]

    return {
        "run_id": run_id,
        "status": "ok",
        "title": req.title,
        "method": summary.get("method"),
        "top_k": summary.get("top_k"),
        "embedding_backend": summary.get("embedding_backend"),
        "claims_evaluated": summary.get("count"),
        "chunk_count": summary.get("chunk_count"),
        "chapter_count": summary.get("chapter_count"),
        "scores": {
            "avg_context_precision": summary.get("avg_context_precision"),
            "avg_context_recall": summary.get("avg_context_recall"),
            "avg_answer_term_coverage": summary.get("avg_answer_score"),
            "avg_latency_seconds": summary.get("avg_latency"),
        },
        "results": evidence,
        "report_url": f"/api/bookproof/report/{run_id}",
    }


# ---------------------------------------------------------------------------
# Public free demo — NO token. 1 validated run per IP per day, strict caps,
# one run at a time, deterministic only, and the run folder (uploaded text +
# results) is deleted immediately after the response is built.
# The token-gated /api/bookproof/v1/verify is unaffected.
# ---------------------------------------------------------------------------


@app.post("/api/bookproof/public/verify", dependencies=[Depends(_single_run_slot)])
async def bookproof_public_verify(
    request: Request,
    book_file: UploadFile = File(...),
    questions_file: Optional[UploadFile] = File(None),
    questions_text: Optional[str] = Form(None),
    method: str = Form(BOOKPROOF_DEFAULT_METHOD),
    top_k: int = Form(5),
) -> Dict[str, Any]:
    if method not in EVAL_METHODS:
        raise HTTPException(status_code=400, detail=f"method must be one of: {', '.join(sorted(EVAL_METHODS))}")
    top_k = max(1, min(int(top_k), 10))

    run_id = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S_") + uuid.uuid4().hex[:8]
    run_dir = RUNS_DIR / ("public_" + run_id)
    run_dir.mkdir(parents=True, exist_ok=True)
    try:
        book_path = await _save_upload(book_file, run_dir, "book_file", BOOKPROOF_BOOK_SUFFIXES)
        if book_path.stat().st_size > PUBLIC_MAX_TEXT_CHARS:
            raise HTTPException(
                status_code=413,
                detail=f"Public demo limit: document must be {PUBLIC_MAX_TEXT_CHARS} characters or fewer. "
                "Use the developer API for larger documents.",
            )

        questions_path = run_dir / "input_questions.jsonl"
        if questions_file is not None and questions_file.filename:
            if Path(_safe_filename(questions_file.filename, "questions.jsonl")).suffix.lower() != ".jsonl":
                raise HTTPException(status_code=400, detail="questions_file must be a .jsonl file")
            raw = (await questions_file.read()).decode("utf-8-sig", errors="replace")
            _validate_questions_jsonl(raw)
            lines = [ln for ln in raw.splitlines() if ln.strip()]
        elif questions_text and questions_text.strip():
            _validate_questions_jsonl(questions_text)
            lines = [ln for ln in questions_text.splitlines() if ln.strip()]
        else:
            raise HTTPException(status_code=400, detail="Provide golden questions as a .jsonl file or as JSONL text")
        if len(lines) > PUBLIC_MAX_CLAIMS:
            raise HTTPException(
                status_code=400,
                detail=f"Public demo limit: max {PUBLIC_MAX_CLAIMS} questions. Use the developer API for more.",
            )
        questions_path.write_text("\n".join(lines) + "\n", encoding="utf-8")

        # Only count a run against the daily quota once it has passed validation.
        _enforce_public_quota(_client_ip(request))

        summary = await run_in_threadpool(
            run_eval_engine,
            book=book_path,
            questions_path=questions_path,
            method=method,
            out=run_dir,
            top_k=top_k,
            embedding_backend="hashing_numpy",
        )

        results: List[Dict[str, Any]] = []
        results_file = run_dir / "results.jsonl"
        if results_file.exists():
            for raw in results_file.read_text(encoding="utf-8").splitlines():
                raw = raw.strip()
                if not raw:
                    continue
                try:
                    results.append(json.loads(raw))
                except json.JSONDecodeError:
                    continue

        preview = [
            {
                "id": row.get("id"),
                "question": " ".join(str(row.get("question") or "").split())[:200],
                "context_precision_like": (row.get("metrics") or {}).get("context_precision_like"),
                "context_recall_like": (row.get("metrics") or {}).get("context_recall_like"),
                "answer_contains_gold_terms": (row.get("metrics") or {}).get("answer_contains_gold_terms"),
            }
            for row in results[:PUBLIC_MAX_CLAIMS]
        ]

        return {
            "run_id": run_id,
            "status": "ok",
            "demo": True,
            "method": summary.get("method"),
            "embedding_backend": summary.get("embedding_backend"),
            "questions_evaluated": summary.get("count"),
            "chunk_count": summary.get("chunk_count"),
            "chapter_count": summary.get("chapter_count"),
            "top_k": summary.get("top_k"),
            "scores": {
                "avg_context_precision": summary.get("avg_context_precision"),
                "avg_context_recall": summary.get("avg_context_recall"),
                "avg_answer_term_coverage": summary.get("avg_answer_score"),
            },
            "results_preview": preview,
            "report_url": None,
            "note": "Free public demo — inputs and results are processed temporarily and deleted immediately. One run per IP per day.",
        }
    finally:
        shutil.rmtree(run_dir, ignore_errors=True)
