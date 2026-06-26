from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Dict, Iterable, List, Tuple


CHAPTER_HEADING_RE = re.compile(
    r"""(?im)^\s{0,8}(?:#{1,6}\s*)?(
        prologue|epilogue|interlude|coda|afterword|
        (?:(?:chapter|chap\.?|book|part|act)\s+[^\n]{1,90})|
        (?:\d{1,4}\.\s+[^\n]{1,90})
    )\s*$""",
    re.VERBOSE,
)

WORD_RE = re.compile(r"\S+")


def _normalize_text(text: str) -> str:
    text = str(text or "").replace("\r\n", "\n").replace("\r", "\n")
    return re.sub(r"\n{3,}", "\n\n", text).strip() + "\n"


def _read_docx_text(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("DOCX input requires python-docx. Run: pip install -r requirements.txt") from exc

    document = Document(str(path))
    parts: List[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return _normalize_text("\n\n".join(parts))


def read_text(path: Path) -> str:
    return read_book_text(path)


def read_book_text(path: Path) -> str:
    path = Path(path)
    suffix = path.suffix.lower()

    if suffix == ".docx":
        return _read_docx_text(path)
    if suffix in {".txt", ""}:
        return _normalize_text(path.read_text(encoding="utf-8-sig"))

    raise ValueError(f"Unsupported book input type: {path.suffix}. Use .txt or .docx")


def word_count(text: str) -> int:
    return len(WORD_RE.findall(text))


def _safe_heading_title(match: re.Match[str]) -> str:
    # Some heading regexes may not define capture group 1; use the whole match safely.
    try:
        raw = match.group(1)
    except IndexError:
        raw = match.group(0)

    if raw is None:
        raw = match.group(0)

    return " ".join(str(raw or "").strip().split()) or "Untitled Chapter"


def detect_chapter_spans(text: str) -> List[Tuple[str, int, int, str]]:
    matches = list(CHAPTER_HEADING_RE.finditer(text))
    if not matches:
        return [("book", 0, len(text), "Book")]

    spans: List[Tuple[str, int, int, str]] = []

    if matches[0].start() > 0:
        preface = text[: matches[0].start()].strip()
        if word_count(preface) >= 80:
            spans.append(("front_matter", 0, matches[0].start(), "front matter"))

    for idx, match in enumerate(matches):
        start = match.start()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
        title = _safe_heading_title(match)
        spans.append((f"ch{idx + 1:03d}", start, end, title))

    return spans


def _split_span_to_chunks(
    text: str,
    span_start: int,
    span_end: int,
    chapter_id: str,
    chapter_title: str,
    chunk_size_words: int,
    overlap_words: int,
) -> List[Dict[str, object]]:
    span_text = text[span_start:span_end]
    words = list(WORD_RE.finditer(span_text))
    if not words:
        return []

    chunk_size_words = max(50, int(chunk_size_words))
    overlap_words = max(0, min(int(overlap_words), chunk_size_words - 1))

    chunks: List[Dict[str, object]] = []
    step = max(1, chunk_size_words - overlap_words)
    word_start = 0
    scene_idx = 1

    while word_start < len(words):
        word_end = min(len(words), word_start + chunk_size_words)
        start_char = span_start + words[word_start].start()
        end_char = span_start + words[word_end - 1].end()
        chunk_text = text[start_char:end_char].strip()
        chunks.append(
            {
                "id": f"{chapter_id}_s{scene_idx:03d}",
                "chapter_id": chapter_id,
                "chapter_title": chapter_title,
                "text": chunk_text,
                "start_char": int(start_char),
                "end_char": int(end_char),
                "word_count": int(word_count(chunk_text)),
            }
        )

        if word_end >= len(words):
            break
        word_start += step
        scene_idx += 1

    return chunks


def chunk_book_text(
    text: str,
    chunk_size_words: int = 900,
    overlap_words: int = 120,
) -> List[Dict[str, object]]:
    chapter_spans = detect_chapter_spans(text)

    if not chapter_spans:
        return _split_span_to_chunks(
            text=text,
            span_start=0,
            span_end=len(text),
            chapter_id="fallback",
            chapter_title="fallback",
            chunk_size_words=chunk_size_words,
            overlap_words=overlap_words,
        )

    chunks: List[Dict[str, object]] = []
    for chapter_id, start, end, title in chapter_spans:
        chapter_chunks = _split_span_to_chunks(
            text=text,
            span_start=start,
            span_end=end,
            chapter_id=chapter_id,
            chapter_title=title,
            chunk_size_words=chunk_size_words,
            overlap_words=overlap_words,
        )
        chunks.extend(chapter_chunks)

    return chunks


def chunk_book_file(
    book_path: Path,
    chunk_size_words: int = 900,
    overlap_words: int = 120,
) -> List[Dict[str, object]]:
    return chunk_book_text(
        read_text(Path(book_path)),
        chunk_size_words=chunk_size_words,
        overlap_words=overlap_words,
    )


def write_jsonl(rows: Iterable[Dict[str, object]], path: Path) -> None:
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="\n") as f:
        for row in rows:
            f.write(json.dumps(row, ensure_ascii=False) + "\n")


def main() -> None:
    parser = argparse.ArgumentParser(description="Chunk a UTF-8 .txt or .docx long-book file.")
    parser.add_argument("--book", required=True, help="Path to a UTF-8 .txt or .docx book file.")
    parser.add_argument("--out", required=True, help="Output chunks JSONL path.")
    parser.add_argument("--chunk-size-words", type=int, default=900)
    parser.add_argument("--overlap-words", type=int, default=120)
    args = parser.parse_args()

    chunks = chunk_book_file(
        Path(args.book),
        chunk_size_words=args.chunk_size_words,
        overlap_words=args.overlap_words,
    )
    write_jsonl(chunks, Path(args.out))
    print(json.dumps({"chunks": len(chunks), "out": args.out}, indent=2))


if __name__ == "__main__":
    main()
